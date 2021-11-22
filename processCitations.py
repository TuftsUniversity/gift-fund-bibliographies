#!/usr/bin/env python3
######################################################################################################
######################################################################################################
######################################################################################################
######################################################################################################
########
########
########    Author:           Henry Steele, Library Technology Services, Tufts University
########    Name of Program:  Input
########	Files:			  processCitations.py, functions.py
########    Date created:     2020-09
########
########    Purpose:
########      - provide Word document format bibliogrphies in Chicago style of books purchased with
########       a lit of gift funds
########
########    Method:
########      - retrieve Analytics reports of MMS IDs and funds for gift funds for given library and year
########        via the Analtyics API
########      - run Alma SRU query for each MMS ID
########      - parse into LaTeX-like output files for rest of processing ( in /Processing.  ".bib" format)
########      - citeproc library (a pseudo-LaTeX Python library) parses into a bibliography and ".docx"
########        library turns them into a Word doc
########
########    Outputs:
########      - A BibTex .bib file containing titles purchased with each fund in the input table (in /Processing)
########      - a Word doc of this data in human readable format suitable for attaching to a donor letter, one for ecah fund
########
########    Dependences:
########      - requirements.txt
########      - install with python3 -m pip install -r requirements.txt
######################################################################################################
######################################################################################################
######################################################################################################
######################################################################################################

import sys
import requests
import json
import os
import csv
import re
import datetime
import codecs
from django.utils.encoding import python_2_unicode_compatible, smart_text, smart_bytes
import io
# from tkinter.filedialog import askopenfilename
import pandas as pd
import numpy as np
import xml.etree.ElementTree as et

import secrets

import docx

import time

sys.path.append('scripts/')
from functions import *




library = input("\n\nWhat library do you want to retreive MMS IDs for? \n\n      1) Tisch\n      2) Ginn\n\nChoose an option: ")

fiscal_year = input("\n\nWhat fiscal year do you want to retreive MMS IDs for? E.g. '2020': ")

parsed_library = ""

if library == "1" or library == "Tisch Library" or library == "Tisch" or library == "TISCH" or library == "tisch" or library == 1:
    parsed_library = "Tisch Library"

elif library == "2" or library == "Ginn Library" or library == "Ginn" or library == "GINN" or library == "ginn" or library == 2:
    parsed_library = "Ginn Library"


# if re.search("^\d{4}$", fiscal_year, re.IGNORECASE) == None:
#     print("Please enter a 4 digit year")
#     sys.exit(1)
# fiscal_year = int(fiscal_year)
# last_year = int(fiscal_year) - 1
# start_date = datetime.date(last_year, 7, 1)
# end_date = datetime.date(fiscal_year, 6, 30)
#
# start_date_string = start_date.strftime("%Y-%m-%d")
# end_date_string = end_date.strftime("%Y-%m-%d")

######################################################################################################
######################################################################################################
#######     composing URL to retrieve Analtyics report, with filter
url = "https://api-na.hosted.exlibrisgroup.com/almaws/v1/analytics/reports?apikey=" + str(secrets.apikey)
limit = "&limit=1000"
format = "&format=xml"


if parsed_library == "Tisch Library":
    path = "&path=%2Fshared%2FTufts%20University%2FReports%2FCollections%2FGift%20Funds%2FTitles%20Purchased%20with%20Gift%20Funds%20-%20Tisch%20-%20Generic%20for%20Script"
elif parsed_library == "Ginn Library":
    path = "&path=%2Fshared%2FTufts%20University%2FReports%2FCollections%2FGift%20Funds%2FTitles%20Purchased%20with%20Gift%20Funds%20-%20Ginn%20-%20Generic%20for%20Script"


#filter = "&filter=%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3Alogical%22+op%3D%22and%22+xmlns%3Asaw%3D%22com.siebel.analytics.web%2Freport%2Fv1.1%22+xmlns%3Asawx%3D%22com.siebel.analytics.web%2Fexpression%2Fv1.1%22+xmlns%3Axsi%3D%22http%3A%2F%2Fwww.w3.org%2F2001%2FXMLSchema-instance%22+xmlns%3Axsd%3D%22http%3A%2F%2Fwww.w3.org%2F2001%2FXMLSchema%22%3E%0D%0A%09%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3Acomparison%22+op%3D%22greaterOrEqual%22%3E%0D%0A%09%09%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3AsqlExpression%22%3E%22Transaction+Date%22.%22Transaction+Date%22%3C%2Fsawx%3Aexpr%3E%0D%0A%09%09%3Csawx%3Aexpr+xsi%3Atype%3D%22xsd%3Adate%22%3E" + start_date_string + "%3C%2Fsawx%3Aexpr%3E%0D%0A%09%3C%2Fsawx%3Aexpr%3E%0D%0A%09%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3Acomparison%22+op%3D%22lessOrEqual%22%3E%0D%0A%09%09%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3AsqlExpression%22%3E%22Transaction+Date%22.%22Transaction+Date%22%3C%2Fsawx%3Aexpr%3E%0D%0A%09%09%3Csawx%3Aexpr+xsi%3Atype%3D%22xsd%3Adate%22%3E" + end_date_string + "%3C%2Fsawx%3Aexpr%3E%0D%0A%09%3C%2Fsawx%3Aexpr%3E%0D%0A%3C%2Fsawx%3Aexpr%3E"

filter = "&filter=%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3Acomparison%22+op%3D%22and%22+xmlns%3Asaw%3D%22com.siebel.analytics.web%2Freport%2Fv1.1%22+%0D%0A++xmlns%3Asawx%3D%22com.siebel.analytics.web%2Fexpression%2Fv1.1%22+%0D%0A++xmlns%3Axsi%3D%22http%3A%2F%2Fwww.w3.org%2F2001%2FXMLSchema-instance%22+%0D%0A++xmlns%3Axsd%3D%22http%3A%2F%2Fwww.w3.org%2F2001%2FXMLSchema%22%3E%0D%0A%09%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3Acomparison%22+op%3D%22notEqual%22%3E%0D%0A%09%09%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3AsqlExpression%22%3E%22Bibliographic+Details%22.%22MMS+Id%22%3C%2Fsawx%3Aexpr%3E%0D%0A%09%09%3Csawx%3Aexpr+xsi%3Atype%3D%22xsd%3Astring%22%3E-1%3C%2Fsawx%3Aexpr%3E%0D%0A%09%3C%2Fsawx%3Aexpr%3E%0D%0A%09%0D%0A%09%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3Acomparison%22+op%3D%22equal%22%3E%0D%0A+++++++++++++++%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3AsqlExpression%22%3E%22Transaction+Date%22.%22Transaction+Date+Fiscal+Year%22%3C%2Fsawx%3Aexpr%3E%3Csawx%3Aexpr+xsi%3Atype%3D%22xsd%3Adecimal%22%3E" + str(fiscal_year) + "%3C%2Fsawx%3Aexpr%3E%0D%0A%09%3C%2Fsawx%3Aexpr%3E%0D%0A%3C%2Fsawx%3Aexpr%3E"



report = requests.get(url + format + path + limit + filter)



report_string = report.content
# #
print(str(report_string))
#
# file = open("test report string.txt", 'w+')
#
# file.write(str(report_string))
# sys.exit()
tree = et.ElementTree(et.fromstring(report.content))

root = tree.getroot()

isFinished = root[0][1].text
resumptionToken = root[0][0].text
isFinishedContinue = isFinished
result_set = root[0][2][0]
##############################################################################################################################
##############################################################################################################################
##############################################################################################################################
########    Retrieve Analytics reports for initial input:
########    - list of MMS IDs from set list of gift funds purchased in the current fiscal year.
########      The list of funds will be be static and part of the Analytics report.
########      The date range will have to be passed as a SAW XML filter to the generic report, for Tisch.
########      The user will set the date range by entering fiscal year as "FY\d\d\d\d" in
########      the first prompt, and choose library by choosing either Tisch or Ginn in the second.
########      The library chosen will affect which report it goes to.
########    - The new version of this report will also contain fund, so that I don't need to retrieve both reports,
########      but for the intiial query it will just use MMS ID
######################################################################################################
######################################################################################################
#######     since, output is limited to 1,000 records, iterate through whole report
while isFinishedContinue == "false":
    full_path = url + "&token=" + resumptionToken
    reportContinue = requests.get(url + "&token=" + resumptionToken)
    # print ("\n\n\n\n" + str(reportContinue.content) + "\n\n\n\n")
    report_string += reportContinue.content

    treeContinue = et.ElementTree(et.fromstring(reportContinue.content))
    rootContinue = treeContinue.getroot()
    result_set_continue = rootContinue[0][1][0]
    result_set.append(result_set_continue)


    isFinishedContinue = rootContinue[0][0].text


mms_id_and_fund_dict = []
mms_id_list = []

x = 0

for element in result_set.iter():


    if re.match(r'.*Row', element.tag):
        mms_id = ""
        fund_code = ""
        fiscal_year = ""
        transacation_item_type = ""

        for subElement in element.iter():

            if re.match(r'.*Column1', subElement.tag):

                mms_id = subElement.text
                mms_id_list.append(mms_id)

            elif re.match(r'.*Column2', subElement.tag):
                fund_code = subElement.text

            elif re.match(r'.*Column5', subElement.tag):
                fiscal_year = subElement.text

            elif re.match(r'.*Column4', subElement.tag):
                transacation_item_type = subElement.text

        mms_id_and_fund_dict.append({'MMS Id': mms_id, 'Fund Ledger Code': fund_code, 'Fiscal Year': fiscal_year, 'Transaction Type': transacation_item_type})

    x += 1
pDir = "./Processing"
if not os.path.isdir(pDir) or not os.path.exists(pDir):
       os.makedirs(pDir)
mms_id_and_fund_df = pd.DataFrame(columns=['MMS Id', 'Fund Ledger Code', 'Fiscal Year', 'Transaction Type'], data=mms_id_and_fund_dict)
mms_id_and_fund_df.to_csv(pDir + "/MMS ID and Fund.csv", index=False)

print(json.dumps(mms_id_and_fund_dict))

y = 0
sru_url = "https://tufts.alma.exlibrisgroup.com/view/sru/01TUN_INST?version=1.2&operation=searchRetrieve&recordSchema=marcxml&query=alma.mms_id="



outfile = open(pDir + "/gifts_funds.txt", "w+", encoding='utf-8')

outfile.write("MMS Id~Author~Author Name~Author Relator~Second Author Name~Second Author Relator~Corporate Author Name~Corporate Author Relator~Second Corporate Author Name~Second Corporate Author Relator~Format~Title~First Place of Publication~First Publisher~First Published Year~Second Place of Publication~Second Publisher~Second Published Year\n")
######################################################################################################
######################################################################################################
#######     iterate through MMS Id list and retrieve bib record by SRU

for id in mms_id_list:

    id = str(id)

    result = requests.get(sru_url + id)




    outfile.write(id + "~")
    print(id)
    tree_bib_record = et.ElementTree(et.fromstring(result.content.decode('utf-8')))
    root_bib_record = tree_bib_record.getroot()

    namespaces = namespaces = {'ns1': 'http://www.loc.gov/MARC21/slim'}

    ######################################################################################################
    ######################################################################################################
    #######     in the returned XML bib record, parse out relevant fields and write to
    #######     tilde-delimited text file
    one_hundred = []
    for record in root_bib_record.findall(".//ns1:datafield[@tag='100']/ns1:subfield[@code='a']", namespaces):
        one_hundred.append(record.text)

    outfile.write(";".join(one_hundred))
    outfile.write("~")


    one_hundred_2 = []
    for record in root_bib_record.findall(".//ns1:datafield[@tag='100']/ns1:subfield[@code='a']", namespaces):
        one_hundred_2.append(record.text)


    try:
        outfile.write(";".join(one_hundred_2))
        outfile.write("~")
    except:
        print(one_hundred_2)
        sys.exit()




    one_hundred_e = []
    for record in root_bib_record.findall(".//ns1:datafield[@tag='100']/ns1:subfield[@code='e']", namespaces):
        one_hundred_e.append(record.text)

    outfile.write(";".join(one_hundred_e))
    outfile.write("~")

    one_ten = []
    for record in root_bib_record.findall(".//ns1:datafield[@tag='110']/ns1:subfield[@code='a']", namespaces):
        one_ten.append(record.text)

    outfile.write(";".join(one_ten))
    outfile.write("~")



    one_ten_e = []
    for record in root_bib_record.findall(".//ns1:datafield[@tag='110']/ns1:subfield[@code='e']", namespaces):
        one_ten_e.append(record.text)

    outfile.write(";".join(one_ten_e))
    outfile.write("~")


    seven_hundred = []
    for record in root_bib_record.findall(".//ns1:datafield[@tag='700']/ns1:subfield[@code='a']", namespaces):
        seven_hundred.append(record.text.encode('utf-8', 'ignore').decode('utf-8', 'ignore'))


    # for element in seven_hundred:
    #     print(id + " - " + str(element.encode('utf-8', 'ignore').decode('utf-8', 'ignore')))
    outfile.write(";".join(seven_hundred).encode('utf-8', 'ignore').decode('utf-8', 'ignore'))
    outfile.write("~")


    seven_hundred_e = []
    for record in root_bib_record.findall(".//ns1:datafield[@tag='700']/ns1:subfield[@code='e']", namespaces):
        seven_hundred_e.append(record.text)

    outfile.write(";".join(seven_hundred_e))
    outfile.write("~")

    seven_ten = []
    for record in root_bib_record.findall(".//ns1:datafield[@tag='710']/ns1:subfield[@code='a']", namespaces):
        seven_ten.append(record.text)

    outfile.write(";".join(seven_ten))
    outfile.write("~")



    seven_ten_e = []
    for record in root_bib_record.findall(".//ns1:datafield[@tag='710']/ns1:subfield[@code='e']", namespaces):
        seven_ten_e.append(record.text)

    outfile.write(";".join(seven_ten_e))
    outfile.write("~")

    six_fifty_five_a = []
    for record in root_bib_record.findall(".//ns1:datafield[@tag='655']/ns1:subfield[@code='a']", namespaces):
        six_fifty_five_a.append(record.text)
    one_sff = ""
    for sff in six_fifty_five_a:
        if 'Electronic' in sff:
            one_sff = sff

    outfile.write(one_sff + "~")

    two_forty_five_a = root_bib_record.find(".//ns1:datafield[@tag='245']/ns1:subfield[@code='a']", namespaces)
    if two_forty_five_a is not None:
        # two_forty_fve_a = re.sub(r'().+?)\s$', r'\1', two_forty_five_a)
        outfile.write(two_forty_five_a.text)


    two_forty_five_b  = root_bib_record.find(".//ns1:datafield[@tag='245']/ns1:subfield[@code='b']", namespaces)
    if two_forty_five_b is not None:
        outfile.write(" ")
        outfile.write(two_forty_five_b.text)

    outfile.write("~")

    two_sixty_a = []
    for record in root_bib_record.findall(".//ns1:datafield[@tag='260']/ns1:subfield[@code='a']", namespaces):
        two_sixty_a.append(record.text)

    outfile.write(";".join(two_sixty_a))
    outfile.write("~")

    two_sixty_b = []
    for record in root_bib_record.findall(".//ns1:datafield[@tag='260']/ns1:subfield[@code='b']", namespaces):
        two_sixty_b.append(record.text)

    outfile.write(";".join(two_sixty_b))
    outfile.write("~")

    two_sixty_c = []
    for record in root_bib_record.findall(".//ns1:datafield[@tag='260']/ns1:subfield[@code='c']", namespaces):
        two_sixty_c.append(record.text)

    outfile.write(";".join(two_sixty_c))
    outfile.write("~")

    two_sixty_four_a = []
    for record in root_bib_record.findall(".//ns1:datafield[@tag='264']/ns1:subfield[@code='a']", namespaces):
        two_sixty_four_a.append(record.text)

    outfile.write(";".join(two_sixty_four_a))
    outfile.write("~")

    two_sixty_four_b = []
    for record in root_bib_record.findall(".//ns1:datafield[@tag='264']/ns1:subfield[@code='b']", namespaces):
        two_sixty_four_b.append(record.text)

    outfile.write(";".join(two_sixty_four_b))
    outfile.write("~")


    two_sixty_four_c = []
    for record in root_bib_record.findall(".//ns1:datafield[@tag='264']/ns1:subfield[@code='c']", namespaces):
        two_sixty_four_c.append(record.text)

    outfile.write(";".join(two_sixty_four_c))
    outfile.write("~")






    outfile.write("\n")

    y += 1

outfile.close()




######################################################################################################
######################################################################################################
#######     method to output bibliography from citeproc
def warn(citation_item):
    print("WARNING: Reference with key '{}' not found in the bibliography."
          .format(citation_item.key))



######################################################################################################
######################################################################################################
#######     retrieve processing files created in previous section
pDir = "./Processing"
marcFilename = "Processing\\gifts_funds.txt"
giftFilename = "Processing\\MMS ID and Fund.csv"


oDir = "./Output"
if not os.path.isdir(oDir) or not os.path.exists(oDir):
       os.makedirs(oDir)


marcDF = pd.read_csv(marcFilename, sep="~", encoding='utf-8', index_col=False)
fundDF = pd.read_csv(giftFilename, encoding='utf-8')

print(marcDF)
print(fundDF)

######################################################################################################
######################################################################################################
#######     make sure all columns are encoded as UTF-8 in dataframes, and merge
for col in marcDF.columns:
    marcDF[col] = marcDF[col].apply(lambda x: smart_bytes(x).decode('utf-8'))

for col in fundDF.columns:
    fundDF[col] = fundDF[col].apply(lambda x: smart_bytes(x).decode('utf-8'))


gf = pd.merge(marcDF, fundDF, on='MMS Id')

print(gf)



for col in gf.columns:

    gf[col] = gf[col].apply(lambda x: x.replace("&", "\&"))

gf= gf.replace('nan', '', regex=True)



gf = gf.drop(gf[gf['Title'].str.isupper()].index)


gf = gf.drop_duplicates(subset=['Title', 'Author'], keep='first')





gf = gf.rename(columns={'Fund Ledger Code':'Fund'})
fundArray = gf.Fund.unique()


fundList = fundArray.tolist()


print(fundList)





######################################################################################################
#######################################################################################################
#######     loop through titles.
#######     for each title, create a bib-formatted
#######     entry in the bib file
#######     receive creator (authors from various)
#######     MARC fields) and relator strings
#######
#######     The string may be a semicolon delimited list, so iterate through
#######     this with a regex match loop and store in an array.  Do the same
#######     for the relators--turn them into an array (list)
#######
#######     Then you can create BibTex statements in the .bib file
#######     for author, translator, and editor


######################################################################################################
######################################################################################################
#######     for each fund, create a folder and .bib file
#######     to store in "/Processing/"
#######

for fund in fundList:

    pFundDir = pDir + "/" + str(fund)
    if not os.path.isdir(pFundDir) or not os.path.exists(pFundDir):
           os.makedirs(pFundDir)





    fundDotBib = str(fund) + ".bib"


    bibFilename = pFundDir + "/" + str(fund) + ".bib"
    outfile = io.open(bibFilename, "w+", encoding='utf-8')

    # outfile.write("\\begin{filecontents}" + str(fund) + ".bib")
    # outfile.write("\n")
    gfSegment = gf.loc[gf['Fund'] == fund]

    gfSegment = gfSegment.reset_index(drop=True)
    print(gfSegment)
    count = len(gfSegment.index)


    print("\nCount: " + str(count))
    x = 0
    while x < count:

        ######################################################################################################
        ######################################################################################################
        #######     for each entry (title record) in the dataframe, separate out and parse metadata into
        #######     BibTex .bib file format
        title = gfSegment.iloc[x]['Title']
        title = re.sub(r'(^.+)\.$', r'\1', title)
        test_mms_id = gfSegment.iloc[x]['MMS Id']
        if title == "":
            x += 1
            continue

        if gfSegment.iloc[x]['Title'].isupper():
            x += 1
            continue
        creator = ""
        if gfSegment.iloc[x]['Author Name'] != "":
            author = gfSegment.iloc[x]['Author Name']

            authorRelator = gfSegment.iloc[x]['Author Relator']
            creator = parseCreator(author, authorRelator, "personal", test_mms_id)
            #creator = parseCreator(author~ authorRelator)
            creator2 = str(author) + "~ " + str(authorRelator) + "\n"
        if gfSegment.iloc[x]['Second Author Name'] != "Empty":
            secondAuthor = gfSegment.iloc[x]['Second Author Name']
            secondAuthorRelator = gfSegment.iloc[x]['Second Author Relator']
            creator += parseCreator(secondAuthor, secondAuthorRelator, "personal", test_mms_id)
            creator2 = str(secondAuthor) + "~ " + str(secondAuthorRelator) + "\n"
        if gfSegment.iloc[x]['Corporate Author Name'] != "Empty":
            corporateAuthor = gfSegment.iloc[x]['Corporate Author Name']
            corporateAuthorRelator = gfSegment.iloc[x]['Corporate Author Relator']
            creator += parseCreator(corporateAuthor, corporateAuthorRelator, "corporate", test_mms_id)
            creator2 = str(corporateAuthor) + "~ " + str(corporateAuthorRelator) + "\n"
        if gfSegment.iloc[x]['Second Corporate Author Name'] != "Empty":
            secondCorporateAuthor = gfSegment.iloc[x]['Second Corporate Author Name']
            secondCorporateAuthorRelator = gfSegment.iloc[x]['Second Corporate Author Relator']
            creator += parseCreator(secondCorporateAuthor, secondCorporateAuthorRelator, "corporate", test_mms_id)
            creator2 = str(secondCorporateAuthor) + "~ " + str(secondCorporateAuthorRelator) + "\n"


        if re.search(r'(author.+?)(\r\n|\r|\n)\t+(author.+?)(\r\n|\r|\n)', creator):
            oneAuthor = re.sub(r'(\tauthor.+?)(\r\n|\r|\n)(\t+author.+?)(\r\n|\r|\n)', r'\1\2', creator)
            creator = oneAuthor
        firstPlacePublished = gfSegment.iloc[x]['First Place of Publication']
        firstPublisher = gfSegment.iloc[x]['First Publisher']
        firstPublishedYear = gfSegment.iloc[x]['First Published Year']
        secondPlacePublished = gfSegment.iloc[x]['Second Place of Publication']
        secondPublisher = gfSegment.iloc[x]['Second Publisher']
        secondPublishedYear = gfSegment.iloc[x]['Second Published Year']

        format = ""
        if re.search(r'[Ee]lectronic', gfSegment.iloc[x]['Format']):
            format = re.sub(r'^.*?([Ee]lectronic\s[A-Za-z- ]+)', r'\1', gfSegment.iloc[x]['Format'])
            print(format)
        publicationInfo = parsePublication(firstPlacePublished, firstPublisher, firstPublishedYear, secondPlacePublished, secondPublisher, secondPublishedYear)



        ######################################################################################################
        ######################################################################################################
        #######     output each entry in the .bib file
        outfile.write("@BOOK{" + gfSegment.iloc[x]['MMS Id'] + ",\n")

        outfile.write(creator)

        if title[-2:] == " /":
            title = title[:-2]
        outfile.write("\ttitle = {" + title + "},\n")
        #
        # if format != "":
        #     publicationInfo = re.sub(r'^(.+?)(\.\")$', r'\1' + r'\1[Online]\2', publicationInfo)
        #     outfile.write("\thowpublished = {" + publicationInfo + "},\n")
        # else:
        outfile.write(publicationInfo)
        #outfile.write(smart_str(creator2))


        if format != "":

            format = re.sub(r'^.*?([Ee]lectronic\s[^; ]+?)', r'. \1', format)
            format = re.sub(r's\.$', '', format)
            print(format)
        note = "\tnote = {<i>" + format + "</i>},\n"
        outfile.write(note)
        outfile.write("}\n\n")

        x += 1

    # outfile.write("\\end{filecontents}")
    # outfile.write("\n")
    #
    # outfile.write("\\documentclass[12pt]{article}")
    # outfile.write("\n")
    # outfile.write("\\usepackage{hyperref}")
    # outfile.write("\n")
    # outfile.write("\\begin{document}")
    # outfile.write("\n")
    #
    # outfile.write("\\cite{test}")
    # outfile.write("\n")
    # outfile.write("\\cite{test1}")
    # outfile.write("\n")
    #
    # outfile.write("\\bibliographystyle{unsrt}")
    # outfile.write("\n")
    # outfile.write("\\bibliography{myRef}")
    # outfile.write("\n")
    #
    # outfile.write("\\end{document}")
    # outfile.write("\n")
    outfile.close()
    from citeproc_local.py2compat import *


    from citeproc_local.source.bibtex import BibTeX

    from citeproc_local import CitationStylesStyle, CitationStylesBibliography
    from citeproc_local import formatter
    from citeproc_local import Citation, CitationItem


    print(bibFilename + "\n")
    file = open(bibFilename, "r+", encoding="utf-8")
    string = file.read()

    print("\n\n" + string + "\n\n" )
    file.close()
    bib_source = BibTeX(bibFilename)


    ######################################################################################################
    ######################################################################################################
    #######     This is the locally included Chicago style template, included in ciceproc/data/styles
    bib_style = CitationStylesStyle('chicago-annotated-bibliography', validate=False)


    # Create the citeproc-py bibliography, passing it the:
    # * CitationStylesStyle,
    # * BibliographySource (BibTeX in this case), and
    # * a formatter (plain, html, or you can write a custom formatter)

    bibliography = CitationStylesBibliography(bib_style, bib_source,
                                              formatter.html)
    ######################################################################################################
    ######################################################################################################
    #######     docx document object
    doc = docx.Document()
    doc.add_heading("References", 0)
    print("References")
    for item in bib_source:

        citation = Citation([CitationItem(item)])



        bibliography.register(citation)
        item_string = bibliography.cite(citation, warn)


    html_helper = HTMLHelper()
    bibliography.sort()
    for item in bibliography.bibliography():
        # print(item)
        # sys.exit()
        ######################################################################################################
        ######################################################################################################
        #######     take out extra characters in citation, that are artifacts of the citeproc citation
        #######     creation process with some of our bib records
        #######
        #######     Also make the editor label plural if there are multiple ("eds.")
        item = str(item)
        item = item.replace(", n.d..", "")
        item = item.replace(',,', ',')
        item = item.replace('..', '.')
        item = re.sub(r'([^<]+?and[^<]+?)(ed.)(\s+<i>)', r'\1eds.\3', item)
        item = item.replace(',.', '.')
        # if re.search(r'^.*?([Ee]lectronic.*?).*?$', item):
            # citation_format = re.sub(r'^.*?\[([Ee]lectronic.*?\]).*?$', r'\1', item)
            # item = item + "\n" + str(citation_format)
            #
            # print(item)
            # # sys.exit()
        ######################################################################################################
        ######################################################################################################
        #######     turn HTML into document styling
        run_map = html_helper.html_to_run_map(str(item))



        par = doc.add_paragraph()

        html_helper.insert_runs_from_html_map(par, run_map)


    doc.save(oDir + "/" + str(fund) + ".docx")
